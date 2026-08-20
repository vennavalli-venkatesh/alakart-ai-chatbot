import os
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from pypdf import PdfReader
import docx


# =========================================
# OCR IMPORTS
# =========================================

_OCR_AVAILABLE = False
_OCR_IMPORT_ERROR = ""

try:
    import pytesseract
    from pdf2image import convert_from_path

    _OCR_AVAILABLE = True

except ImportError as e:
    _OCR_IMPORT_ERROR = str(e)


# =========================================
# PDF TEXT THRESHOLD
# =========================================

MIN_TEXT_THRESHOLD = 50


# =========================================
# DOCUMENT MODEL
# =========================================

@dataclass
class Document:
    """
    Represents one loaded document/page/chunk
    with its text and metadata.
    """

    page_content: str
    metadata: Dict[str, Any]


# =========================================
# DOCUMENT LOADER
# =========================================

class DocumentLoader:
    """
    Loads all knowledge-base documents.

    The loader strictly separates:

        products   -> Alakart product catalogue
        otc        -> General medicine / OTC information
        wellness   -> General wellness information
        navigation -> Alakart app/navigation information
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".json",
        ".txt",
        ".md",
    }

    VALID_CATEGORIES = {
        "products",
        "otc",
        "wellness",
        "navigation",
    }

    def __init__(self, base_data_dir: str = "data"):
        self.base_data_dir = Path(base_data_dir)

    # =========================================
    # LOAD SINGLE FILE
    # =========================================

    def load_file(
        self,
        file_path: str,
        category: Optional[str] = None,
    ) -> List[Document]:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            print(
                f"Skipping unsupported file extension: {file_path}"
            )
            return []

        # If category is not explicitly provided,
        # use the parent folder name.
        if category is None:
            category = path.parent.name

        category = str(category).strip().lower()

        # Prevent accidental unknown categories.
        if category not in self.VALID_CATEGORIES:
            print(
                f"WARNING: Unknown knowledge category "
                f"'{category}' for file '{path.name}'."
            )

        # -----------------------------------------
        # Build metadata
        # -----------------------------------------

        base_metadata = self._build_metadata(
            path=path,
            category=category,
        )

        # -----------------------------------------
        # Load according to file type
        # -----------------------------------------

        if ext == ".pdf":
            return self._load_pdf(
                path,
                base_metadata,
            )

        if ext == ".docx":
            return self._load_docx(
                path,
                base_metadata,
            )

        if ext == ".json":
            return self._load_json(
                path,
                base_metadata,
            )

        if ext in {".txt", ".md"}:
            return self._load_text(
                path,
                base_metadata,
            )

        return []

    # =========================================
    # METADATA BUILDER
    # =========================================

    def _build_metadata(
        self,
        path: Path,
        category: str,
    ) -> Dict[str, Any]:

        category = str(category).strip().lower()

        # -----------------------------------------
        # Default values
        # -----------------------------------------

        source_type = "general_knowledge"
        document_type = "general_knowledge"

        # -----------------------------------------
        # PRODUCTS
        # -----------------------------------------

        if category == "products":

            source_type = "alakart_catalogue"
            document_type = "product_catalog"

        # -----------------------------------------
        # OTC
        # -----------------------------------------

        elif category == "otc":

            source_type = "otc_reference"
            document_type = "otc_information"

        # -----------------------------------------
        # WELLNESS
        # -----------------------------------------

        elif category == "wellness":

            source_type = "wellness_reference"
            document_type = "wellness_information"

        # -----------------------------------------
        # NAVIGATION
        # -----------------------------------------

        elif category == "navigation":

            source_type = "navigation_reference"
            document_type = "navigation"

        # -----------------------------------------
        # Final metadata
        # -----------------------------------------

        return {
            # Main classification
            "category": category,

            # Human-readable source
            "source": path.name,
            "source_name": path.name,

            # Machine-readable classification
            "source_type": source_type,
            "document_type": document_type,

            # File information
            "file_path": str(path),
            "file_type": path.suffix.lower().lstrip("."),
        }

    # =========================================
    # PDF LOADER
    # =========================================

    def _load_pdf(
        self,
        path: Path,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:

        documents: List[Document] = []

        try:

            reader = PdfReader(str(path))
            total_pages = len(reader.pages)

            print("\n" + "=" * 60)
            print(f"PDF: {path.name}")
            print(f"Total pages: {total_pages}")
            print("=" * 60)

            for index, page in enumerate(reader.pages):

                page_number = index + 1

                # -----------------------------------------
                # Normal extraction
                # -----------------------------------------

                normal_text = page.extract_text() or ""
                normal_text = normal_text.strip()

                normal_length = len(normal_text)

                final_text = normal_text
                ocr_used = False

                # -----------------------------------------
                # OCR fallback
                # -----------------------------------------

                if normal_length < MIN_TEXT_THRESHOLD:

                    print(
                        f"Page {page_number}: "
                        f"normal text length = {normal_length}. "
                        f"Trying OCR..."
                    )

                    if _OCR_AVAILABLE:

                        ocr_used = True

                        try:

                            ocr_text = self._ocr_page(
                                pdf_path=str(path),
                                page_number=index,
                            )

                            if ocr_text.strip():
                                final_text = ocr_text.strip()

                        except Exception as e:

                            print(
                                f"WARNING: OCR failed on "
                                f"{path.name}, page "
                                f"{page_number}: {e}"
                            )

                    else:

                        print(
                            f"WARNING: OCR unavailable for "
                            f"{path.name}, page "
                            f"{page_number}."
                        )

                # -----------------------------------------
                # Logging
                # -----------------------------------------

                print(
                    f"Page {page_number}: "
                    f"normal={normal_length}, "
                    f"final={len(final_text)}, "
                    f"OCR={'YES' if ocr_used else 'NO'}"
                )

                # -----------------------------------------
                # Create Document
                # -----------------------------------------

                if final_text.strip():

                    metadata = base_metadata.copy()

                    metadata.update({
                        "page": page_number,
                        "total_pages": total_pages,
                        "ocr_used": ocr_used,
                        "text_length": len(final_text),
                    })

                    documents.append(
                        Document(
                            page_content=final_text,
                            metadata=metadata,
                        )
                    )

                else:

                    print(
                        f"WARNING: Page {page_number} "
                        f"produced no usable text."
                    )

        except Exception as e:

            print(
                f"ERROR loading PDF '{path}': {e}"
            )

        return documents

    # =========================================
    # OCR SINGLE PAGE
    # =========================================

    def _ocr_page(
        self,
        pdf_path: str,
        page_number: int,
    ) -> str:

        if not _OCR_AVAILABLE:

            raise RuntimeError(
                "OCR dependencies are not installed.\n"
                f"Import error: {_OCR_IMPORT_ERROR}\n\n"
                "Install:\n"
                "pip install pytesseract pdf2image Pillow\n\n"
                "Also install Tesseract OCR and Poppler."
            )

        images = convert_from_path(
            pdf_path,
            first_page=page_number + 1,
            last_page=page_number + 1,
            dpi=300,
        )

        if not images:
            return ""

        text = pytesseract.image_to_string(
            images[0],
            lang="eng",
        )

        return text

    # =========================================
    # DOCX LOADER
    # =========================================

    def _load_docx(
        self,
        path: Path,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:

        documents: List[Document] = []

        try:

            doc = docx.Document(str(path))

            # -----------------------------------------
            # Paragraphs
            # -----------------------------------------

            paragraphs = [
                paragraph.text.strip()
                for paragraph in doc.paragraphs
                if paragraph.text.strip()
            ]

            # -----------------------------------------
            # Tables
            # -----------------------------------------

            table_lines: List[str] = []

            for table in doc.tables:

                for row in table.rows:

                    row_values = []

                    for cell in row.cells:

                        value = (
                            cell.text
                            .strip()
                            .replace("\n", " ")
                        )

                        row_values.append(value)

                    row_text = " | ".join(row_values)

                    if row_text.strip():
                        table_lines.append(row_text)

            # -----------------------------------------
            # Combine content
            # -----------------------------------------

            sections = []

            if paragraphs:
                sections.append(
                    "\n".join(paragraphs)
                )

            if table_lines:
                sections.append(
                    "--- Tables ---\n"
                    + "\n".join(table_lines)
                )

            full_content = "\n\n".join(sections).strip()

            # -----------------------------------------
            # Metadata
            # -----------------------------------------

            metadata = base_metadata.copy()

            metadata.update({
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "text_length": len(full_content),
            })

            # -----------------------------------------
            # Create Document
            # -----------------------------------------

            if full_content:

                documents.append(
                    Document(
                        page_content=full_content,
                        metadata=metadata,
                    )
                )

            else:

                print(
                    f"WARNING: DOCX contains no "
                    f"extractable text: {path}"
                )

        except Exception as e:

            print(
                f"ERROR loading DOCX '{path}': {e}"
            )

        return documents

    # =========================================
    # JSON LOADER
    # =========================================

    def _load_json(
        self,
        path: Path,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:

        documents: List[Document] = []

        try:

            with open(
                path,
                "r",
                encoding="utf-8",
            ) as file:

                data = json.load(file)

            # -----------------------------------------
            # Validate product JSON structure
            # -----------------------------------------

            if (
                base_metadata.get("category") == "products"
                and not isinstance(data, list)
            ):

                print(
                    f"WARNING: Product JSON '{path.name}' "
                    f"is not a JSON array."
                )

            # -----------------------------------------
            # Convert JSON to readable text
            # -----------------------------------------

            content = json.dumps(
                data,
                indent=2,
                ensure_ascii=False,
            )

            # -----------------------------------------
            # Metadata
            # -----------------------------------------

            metadata = base_metadata.copy()

            metadata["text_length"] = len(content)

            if isinstance(data, list):
                metadata["record_count"] = len(data)

            # -----------------------------------------
            # Create Document
            # -----------------------------------------

            if content.strip():

                documents.append(
                    Document(
                        page_content=content,
                        metadata=metadata,
                    )
                )

        except json.JSONDecodeError as e:

            print(
                f"ERROR: Invalid JSON in '{path}': {e}"
            )

        except Exception as e:

            print(
                f"ERROR loading JSON '{path}': {e}"
            )

        return documents

    # =========================================
    # TXT / MARKDOWN LOADER
    # =========================================

    def _load_text(
        self,
        path: Path,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:

        documents: List[Document] = []

        try:

            with open(
                path,
                "r",
                encoding="utf-8",
            ) as file:

                text = file.read().strip()

            if not text:
                return documents

            metadata = base_metadata.copy()

            metadata["text_length"] = len(text)

            documents.append(
                Document(
                    page_content=text,
                    metadata=metadata,
                )
            )

        except Exception as e:

            print(
                f"ERROR loading text file "
                f"'{path}': {e}"
            )

        return documents

    # =========================================
    # DIRECTORY LOADING
    # =========================================

    def load_directory(
        self,
        dir_path: str,
        category: Optional[str] = None,
    ) -> List[Document]:

        directory = Path(dir_path)

        if not directory.exists():

            print(
                f"Directory not found: {dir_path}"
            )

            return []

        if not directory.is_dir():

            print(
                f"Not a directory: {dir_path}"
            )

            return []

        # -----------------------------------------
        # Determine category ONCE
        # -----------------------------------------

        final_category = (
            str(category).strip().lower()
            if category
            else directory.name.lower()
        )

        documents: List[Document] = []

        # -----------------------------------------
        # Walk directory
        # -----------------------------------------

        for root, _, files in os.walk(directory):

            for filename in sorted(files):

                file_path = Path(root) / filename

                if (
                    file_path.suffix.lower()
                    not in self.SUPPORTED_EXTENSIONS
                ):
                    continue

                try:

                    loaded_documents = self.load_file(
                        str(file_path),
                        category=final_category,
                    )

                    documents.extend(
                        loaded_documents
                    )

                except Exception as e:

                    print(
                        f"ERROR loading "
                        f"'{file_path}': {e}"
                    )

        return documents

    # =========================================
    # LOAD ALL KNOWLEDGE SOURCES
    # =========================================

    def load_all_knowledge_sources(
        self,
        categories: Optional[List[str]] = None,
    ) -> List[Document]:

        if categories is None:

            categories = [
                "navigation",
                "otc",
                "products",
                "wellness",
            ]

        all_documents: List[Document] = []

        # -----------------------------------------
        # Load each category separately
        # -----------------------------------------

        for category in categories:

            category = (
                str(category)
                .strip()
                .lower()
            )

            if category not in self.VALID_CATEGORIES:

                print(
                    f"WARNING: Unknown category "
                    f"'{category}'. Skipping."
                )

                continue

            category_directory = (
                self.base_data_dir / category
            )

            if not category_directory.exists():

                print(
                    f"Knowledge directory does not exist: "
                    f"{category_directory}"
                )

                continue

            documents = self.load_directory(
                str(category_directory),
                category=category,
            )

            print(
                f"Loaded {len(documents)} documents "
                f"from category: {category}"
            )

            all_documents.extend(documents)

        # -----------------------------------------
        # Final statistics
        # -----------------------------------------

        print(
            "\nTotal loaded documents: "
            f"{len(all_documents)}"
        )

        return all_documents